"""
modules/docx_exporter.py — professional layout 
Includes soil EMBRAPA section and fixes hist_baseline parameter.
"""
from __future__ import annotations
from datetime import date, datetime
from typing import Dict, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import CROP_PARAMS, CLIMATE_NORMALS_RS, SOIL_APTITUDE_CLASSES

#color palette
C_AZUL_ESC  = RGBColor(0x1A, 0x5C, 0x96)
C_AZUL_MED  = RGBColor(0x27, 0x80, 0xC8)
C_VERDE     = RGBColor(0x1E, 0x8B, 0x4C)
C_VERMELHO  = RGBColor(0xC0, 0x39, 0x2B)
C_AMARELO   = RGBColor(0xB7, 0x77, 0x00)
C_CINZA     = RGBColor(0x55, 0x55, 0x55)
C_CINZA_CLR = RGBColor(0x88, 0x88, 0x88)
C_BRANCO    = RGBColor(0xFF, 0xFF, 0xFF)
C_VERDE_ESC = RGBColor(0x15, 0x6B, 0x38)
C_MARROM    = RGBColor(0x7B, 0x4B, 0x1A)

HEX_AZUL_ESC = "1A5C96"
HEX_AZUL_HDR = "2780C8"
HEX_ROW_ALT  = "EBF3FB"
HEX_VERDE_BG = "E8F5EE"
HEX_VERM_BG  = "FDECEA"
HEX_AMAR_BG  = "FFF8E1"
HEX_SOLO_BG  = "F5EFE6"
HEX_SOLO_HDR = "8B6340"

EVENT_LABELS = {
    "drought":  "Drought / Water Deficit",
    "rainfall": "Excess Rainfall / Flooding",
    "frost":    "Frost",
    "hail":     "Hail / Severe Storm",
}
INDEX_META = {
    "NDVI": "Vegetation Index (NDVI)",
    "NDRE": "NDVI Red-Edge — early stress",
    "EVI":  "Enhanced Vegetation Index (EVI)",
    "NDWI": "Canopy Water Index (NDWI)",
    "NDMI": "SWIR Moisture Index (NDMI)",
    "BSI":  "Bare Soil Index (BSI)",
    "NBR":  "Normalised Burn Ratio (NBR)",
    "PSRI": "Senescence Index (PSRI)",
    "CRI1": "Carotenoid Index (CRI1)",
    "VHI":  "Vegetation Health Index (VHI)",
}


#formatting utilities
def _brl(v):
    return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def _pct(v):
    return f"{v:.1f}%"

def _shd(cell_or_para, hex_c):
    obj = cell_or_para._tc if hasattr(cell_or_para, "_tc") else cell_or_para._p
    pr  = (obj.get_or_add_tcPr() if hasattr(cell_or_para, "_tc")
           else obj.get_or_add_pPr())
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hex_c)
    pr.append(s)

def _cell_margins(cell, t=80, b=80, l=120, r=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in [("top", t), ("bottom", b), ("left", l), ("right", r)]:
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)

def _para_bot_border(para, color="2780C8", sz="8"):
    pPr = para._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), sz)
    bot.set(qn("w:space"), "4")
    bot.set(qn("w:color"), color)
    pb.append(bot)
    pPr.append(pb)

def _run(para, text, bold=False, italic=False, size=10, color=None):
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.name = "Arial"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color

def _col_widths(table, cms):
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    total = int(sum(cms) * 567)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:w"), str(total))
    tw.set(qn("w:type"), "dxa")
    ex = tblPr.find(qn("w:tblW"))
    if ex is not None:
        tblPr.remove(ex)
    tblPr.append(tw)
    grid = OxmlElement("w:tblGrid")
    for w in cms:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(w * 567)))
        grid.append(gc)
    eg = tbl.find(qn("w:tblGrid"))
    if eg is not None:
        tbl.remove(eg)
    tbl.insert(1, grid)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tw2  = OxmlElement("w:tcW")
            tw2.set(qn("w:w"), str(int(cms[i] * 567)))
            tw2.set(qn("w:type"), "dxa")
            ex2 = tcPr.find(qn("w:tcW"))
            if ex2 is not None:
                tcPr.remove(ex2)
            tcPr.append(tw2)
            _cell_margins(cell)

def _hdr_row(table, headers, bg=HEX_AZUL_HDR):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        _shd(cell, bg)
        _cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, bold=True, size=9, color=C_BRANCO)

def _data_row(table, values, alt=False, colors=None):
    row = table.add_row()
    if alt:
        for c in row.cells:
            _shd(c, HEX_ROW_ALT)
    for i, v in enumerate(values):
        cell = row.cells[i]
        _cell_margins(cell)
        p    = cell.paragraphs[0]
        clr  = colors[i] if colors and i < len(colors) else None
        _run(p, str(v), size=9, color=clr, bold=(clr is not None))
    return row

def _add_footer(doc):
    sec  = doc.sections[0]
    ftr  = sec.footer
    para = ftr.paragraphs[0]
    para.clear()
    para.paragraph_format.space_before = Pt(4)
    _run(para, "Poseidon-Copernicus-EMBRAPA Validator  —  Agricultural Insurance Claim Report",
         size=8, color=C_CINZA_CLR, italic=True)
    para.add_run("\t")
    r = para.add_run()
    r.font.size = Pt(8)
    r.font.color.rgb = C_CINZA
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r._r.append(fc_begin)
    ins = OxmlElement("w:instrText")
    ins.set(qn("xml:space"), "preserve")
    ins.text = " PAGE "
    r._r.append(ins)
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r._r.append(fc_end)
    _para_bot_border(para)
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    for pos, val in [("3780", "center"), ("7560", "right")]:
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), val)
        tab.set(qn("w:pos"), pos)
        tabs.append(tab)
    pPr.insert(0, tabs)


#main class
class DocxExporter:

    def __init__(
        self, event_type, crop_type, start_date, end_date, area_ha,
        farm_name="Rural Property", planting_date=None, centroid=None,
    ):
        self.event_type    = event_type
        self.crop_type     = crop_type
        self.start_date    = start_date
        self.end_date      = end_date
        self.area_ha       = area_ha
        self.farm_name     = farm_name
        self.planting_date = planting_date
        self.centroid      = centroid or {}
        self.crop_params   = CROP_PARAMS.get(crop_type, CROP_PARAMS["soybean"])
        self.doc           = Document()
        for sec in self.doc.sections:
            sec.top_margin    = Cm(2.0)
            sec.bottom_margin = Cm(2.0)
            sec.left_margin   = Cm(2.5)
            sec.right_margin  = Cm(2.5)
        self.doc.styles["Normal"].font.name = "Arial"
        self.doc.styles["Normal"].font.size = Pt(10)
        _add_footer(self.doc)

    #fixed entry point
    def export(
        self,
        analysis:      Dict,
        cop_data:      Dict,
        pos_summ:      Dict,
        pos_vote:      Dict,
        output_path:   str,
        hist_baseline: Dict = None,   #was missing
        soil_data:     Dict = None,   #soil data
    ) -> str:
        self._cover()

        self._sec("1. CLAIM IDENTIFICATION")
        self._ident()

        self._sec("2. ECONOMIC CONTEXT")
        self._context()

        self._sec("3. SATELLITE DATA — COPERNICUS / SENTINEL-2")
        self._satellite(cop_data)

        self._sec("4. WEATHER DATA — POSEIDON NETWORK")
        self._poseidon(pos_summ, pos_vote)

        #soil section — only displays if data is available.
        if soil_data and not soil_data.get("error"):
            self._sec("5. SOIL ANALYSIS — EMBRAPA / AGRICULTURAL SUITABILITY")
            self._soil(soil_data, analysis)
            self._sec("6. VALIDATION CHECKLIST")
            self._checks(analysis)
            self._sec("7. ESTIMATED ECONOMIC LOSSES")
            self._loss(analysis, hist_baseline or {})
        else:
            self._sec("5. VALIDATION CHECKLIST")
            self._checks(analysis)
            self._sec("6. ESTIMATED ECONOMIC LOSSES")
            self._loss(analysis, hist_baseline or {})

        self._verdict(analysis)
        self._note(soil_data)
        self.doc.save(output_path)
        return output_path

    #helper section
    def _sec(self, title):
        doc = self.doc
        doc.add_paragraph()
        p = doc.add_paragraph()
        _shd(p, HEX_AZUL_ESC)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(6)
        _run(p, f"  {title}", bold=True, size=11, color=C_BRANCO)

    #cover
    def _cover(self):
        doc = self.doc
        for _ in range(4):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, "POSEIDON  ✦  COPERNICUS  ✦  SENTINEL-2  ✦  EMBRAPA",
             bold=True, size=11, color=C_AZUL_MED)
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shd(p, HEX_AZUL_ESC)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(10)
        _run(p, "AGRICULTURAL INSURANCE CLAIM VALIDATION REPORT",
             bold=True, size=16, color=C_BRANCO)
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, self.farm_name, bold=True, size=15, color=C_AZUL_ESC)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, self.crop_params["name_en"], size=11, color=C_CINZA)
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _shd(p, HEX_VERM_BG)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        _run(p, EVENT_LABELS.get(self.event_type, self.event_type).upper(),
             bold=True, size=13, color=C_VERMELHO)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, (f"{self.start_date.strftime('%d/%m/%Y')} a {self.end_date.strftime('%d/%m/%Y')}"
                 f"  |  {self.area_ha:.0f} hectares"), size=11, color=C_CINZA)
        for _ in range(3):
            doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_bot_border(p, color=HEX_AZUL_HDR, sz="12")
        _run(p, f"Generated on {datetime.now().strftime('%d/%m/%Y at %H:%M')}",
             size=9, color=C_CINZA_CLR, italic=True)
        doc.add_page_break()

    #identification
    def _ident(self):
        doc = self.doc
        doc.add_paragraph()
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Field", "Information"])
        _col_widths(tbl, [5.0, 10.5])
        fields = [
            ("Propriedade",   self.farm_name),
            ("Crop",       self.crop_params["name_en"]),
            ("Alleged Event", EVENT_LABELS.get(self.event_type, self.event_type)),
            ("Period",       f"{self.start_date.strftime('%d/%m/%Y')} a {self.end_date.strftime('%d/%m/%Y')}"),
            ("Duration",       f"{(self.end_date - self.start_date).days + 1} days"),
            ("Area",          f"{self.area_ha:.1f} hectares"),
            ("Centroid",     f"Lat {self.centroid.get('lat','N/D')} | Lon {self.centroid.get('lon','N/D')}"),
            ("Planting Date", self.planting_date.strftime("%d/%m/%Y") if self.planting_date else "Not provided"),
        ]
        for i, (k, v) in enumerate(fields):
            row = tbl.add_row()
            alt = (i % 2 == 1)
            for j, (cell, txt) in enumerate(zip(row.cells, [k, v])):
                if alt:
                    _shd(cell, HEX_ROW_ALT)
                _cell_margins(cell)
                p = cell.paragraphs[0]
                _run(p, txt, bold=(j == 0), size=10)

    #context
    def _context(self):
        doc = self.doc
        cp  = self.crop_params
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.8)
        _run(p, (
            f"{cp['name_en']} crop with average historical yield of "
            f"{cp['yield_avg_bags_ha']} bags/ha (range: {cp['yield_min_bags_ha']}–"
            f"{cp['yield_max_bags_ha']} bags/ha). For {self.area_ha:.0f} ha at "
            f"{_brl(cp['price_brl_bag'])}/bag, expected gross revenue is "
            f"{_brl(cp['yield_avg_bags_ha'] * self.area_ha * cp['price_brl_bag'])} (ref. CEPEA)."
        ), size=10)

    #satellite
    def _satellite(self, cop):
        doc = self.doc
        doc.add_paragraph()
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Index", "Description", "Baseline", "Event", "Δ%", "Status"])
        _col_widths(tbl, [1.4, 5.8, 2.0, 2.0, 1.6, 2.7])
        for i, (idx, desc) in enumerate(INDEX_META.items()):
            if idx not in cop or "error" in cop[idx]:
                continue
            d    = cop[idx]
            b    = d.get("baseline_mean")
            e    = d.get("event_mean")
            apct = d.get("anomaly_pct")
            b_s  = f"{b:.4f}" if b is not None else "N/D"
            e_s  = f"{e:.4f}" if e is not None else "N/D"
            p_s  = f"{apct:+.1f}%" if apct is not None else "N/D"
            if apct is not None:
                drop = -apct
                if   drop >= 20: status, sc = "CRITICAL ▼", C_VERMELHO
                elif drop >= 10: status, sc = "LOW ↓",   C_AMARELO
                else:            status, sc = "NORMAL",    C_VERDE
            else:
                status, sc = "N/D", C_CINZA
            _data_row(tbl, [idx, desc, b_s, e_s, p_s, status],
                      alt=(i % 2 == 1), colors=[None, None, None, None, None, sc])
        doc.add_paragraph()
        if "VHI" in cop and cop["VHI"].get("event_mean") is not None:
            vhi = cop["VHI"]["event_mean"]
            vci = cop["VHI"].get("vci", "N/D")
            tci = cop["VHI"].get("tci", "N/D")
            if   vhi < 35: vc, lbl, bg = C_VERMELHO, "ESTRESSE SEVERO",    HEX_VERM_BG
            elif vhi < 50: vc, lbl, bg = C_AMARELO,  "ESTRESSE MODERADO",  HEX_AMAR_BG
            else:          vc, lbl, bg = C_VERDE,    "HEALTHY VEGETATION", HEX_VERDE_BG
            p = doc.add_paragraph()
            _shd(p, bg)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            _run(p, f"  VHI: {vhi:.1f} — {lbl}  |  VCI: {vci}  |  TCI: {tci}",
                 bold=True, size=10, color=vc)
        doc.add_paragraph()

    #poseidon
    def _poseidon(self, summary, vote):
        doc = self.doc
        doc.add_paragraph()
        if summary:
            mid     = self.start_date.month
            normals = CLIMATE_NORMALS_RS.get(mid, {})
            pd2     = summary.get("period_days", 30)
            months  = pd2 / 30
            np_     = normals.get("prcp_mm", 110) * months
            pt      = summary.get("prcp_total_mm", 0)
            pp      = pt / np_ * 100 if np_ else 0
            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"
            _hdr_row(tbl, ["Weather Variable", "Measured Value", "Reference / Context"])
            _col_widths(tbl, [6.0, 3.5, 6.0])
            rows_d = [
                ("Cumulative precipitation",    f"{pt:.1f} mm",                            f"{pp:.0f}% of normal ({np_:.0f} mm)"),
                ("Average temperature",         f"{summary.get('tavg_mean_c',0):.2f} °C",  f"Normal: {normals.get('tavg_c','N/D')} °C"),
                ("Abs. max. temperature",   f"{summary.get('tmax_abs_c',0):.2f} °C",   "—"),
                ("Abs. min. temperature",   f"{summary.get('tmin_abs_c',0):.2f} °C",   "—"),
                ("Average relative humidity",    f"{summary.get('rh_avg_mean_pct',0):.1f}%", "—"),
                ("Rainy days >1 mm",      str(summary.get("prcp_days", "N/D")),       "—"),
                ("Max. wind speed",              f"{summary.get('wspd_max_kmh',0):.1f} km/h","—"),
            ]
            for i, (k, v, ctx) in enumerate(rows_d):
                row = tbl.add_row()
                alt = (i % 2 == 1)
                for j, (cell, txt) in enumerate(zip(row.cells, [k, v, ctx])):
                    if alt:
                        _shd(cell, HEX_ROW_ALT)
                    _cell_margins(cell)
                    p = cell.paragraphs[0]
                    _run(p, txt, bold=(j == 0), size=10)
            doc.add_paragraph()

        ws     = vote.get("weighted_score", 0.0)
        sl     = vote.get("signal_level", "desconhecido")
        passed = vote.get("passed", False)
        votes  = vote.get("votes", {})
        bg     = HEX_VERDE_BG if passed else HEX_VERM_BG
        cor    = C_VERDE if passed else C_VERMELHO
        ico    = "✔ APPROVED" if passed else "✘ REJECTED"
        p = doc.add_paragraph()
        _shd(p, bg)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        _run(p, f"  Poseidon IDW Climate Score: {ws:.0f}/100 — signal {sl}  |  {ico}",
             bold=True, size=10, color=cor)
        if votes:
            doc.add_paragraph()
            tbl2 = doc.add_table(rows=1, cols=4)
            tbl2.style = "Table Grid"
            _hdr_row(tbl2, ["Direction", "Point ID", "Result", "Details"])
            _col_widths(tbl2, [2.0, 2.0, 3.0, 8.5])
            dl = {"N": "North ↑", "S": "South ↓", "E": "East →", "W": "West ←"}
            for i, (direction, v) in enumerate(votes.items()):
                ok = v.get("confirmed", False)
                c2 = C_VERDE if ok else C_VERMELHO
                _data_row(tbl2, [
                    dl.get(direction, direction),
                    str(v.get("point_id", "N/D")),
                    "CONFIRMED ✔" if ok else "NOT CONFIRMED ✘",
                    v.get("reason", "")[:90],
                ], alt=(i % 2 == 1), colors=[None, None, c2, None])
        doc.add_paragraph()

    #soil EMBRAPA
    def _soil(self, soil_data: Dict, analysis: Dict):
        doc = self.doc
        doc.add_paragraph()

        dominant_class = soil_data.get("dominant_class")
        soil_name      = soil_data.get("resolved_name") or soil_data.get("soil_name", "N/D")
        soil_code      = soil_data.get("soil_code", "N/D")
        suitable       = soil_data.get("suitable_for_agriculture", True)
        apt_label      = soil_data.get("aptitude_label", "N/D")
        apt_desc       = soil_data.get("aptitude_description", "")
        dom_pct        = soil_data.get("dominant_percentage", 0)
        water          = soil_data.get("water_props", {})
        cl_pct         = soil_data.get("classified_area_percentage", 0)

        #fitness table
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Soil Attribute", "Value / Description"], bg=HEX_SOLO_HDR)
        _col_widths(tbl, [5.5, 10.0])
        fields = [
            ("EMBRAPA Suitability Class",   f"Class {dominant_class} — {apt_label}"),
            ("Crop Suitability",        "SUITABLE ✔" if suitable else "UNSUITABLE ✘"),
            ("Dominant Soil",              f"{soil_code} | {soil_name}"),
            ("Area Share",         f"{dom_pct:.0f}% of field area"),
            ("Classified Area",            f"{cl_pct:.0f}%"),
            ("Description",                    apt_desc),
        ]
        for i, (k, v) in enumerate(fields):
            row = tbl.add_row()
            alt = (i % 2 == 1)
            cor_v = C_VERDE if (k == "Crop Suitability" and suitable) else \
                    C_VERMELHO if (k == "Crop Suitability" and not suitable) else None
            for j, (cell, txt) in enumerate(zip(row.cells, [k, v])):
                _shd(cell, HEX_SOLO_BG if alt else "FFFFFF")
                _cell_margins(cell)
                p = cell.paragraphs[0]
                _run(p, txt, bold=(j == 0), size=10, color=(cor_v if j == 1 else None))
        doc.add_paragraph()

        #water properties
        if water:
            awc = water.get("AWC", "N/D")
            ks  = water.get("Ks",  "N/D")
            fc  = water.get("fc",  "N/D")
            wp  = water.get("wp",  "N/D")
            ret = water.get("retention", "N/D")
            tex = water.get("texture",  "N/D")

            tbl2 = doc.add_table(rows=1, cols=3)
            tbl2.style = "Table Grid"
            _hdr_row(tbl2, ["Hydraulic Property", "Value", "Interpretation"], bg=HEX_SOLO_HDR)
            _col_widths(tbl2, [5.5, 3.0, 7.0])
            hid_rows = [
                ("AWC — Plant-Available Water",    f"{awc} mm/m",  "Water storage capacity available to plants"),
                ("Ks — Hydraulic Conductivity", f"{ks} mm/h",   "Saturated drainage speed"),
                ("Field Capacity",      f"{fc}%",       "Water content retained after free drainage"),
                ("Wilting Point",          f"{wp}%",       "Minimum water content for plant survival"),
                ("Water Retention",         ret,            "General retention classification"),
                ("Texture",                  tex,            "Granulometria dominante do solo"),
            ]
            for i, (k, v, ctx) in enumerate(hid_rows):
                row = tbl2.add_row()
                alt = (i % 2 == 1)
                for j, (cell, txt) in enumerate(zip(row.cells, [k, v, ctx])):
                    _shd(cell, HEX_SOLO_BG if alt else "FFFFFF")
                    _cell_margins(cell)
                    p = cell.paragraphs[0]
                    _run(p, txt, bold=(j == 0), size=10)
            doc.add_paragraph()

        #other soils in the plot
        soil_types = soil_data.get("soil_types", [])
        if len(soil_types) > 1:
            tbl3 = doc.add_table(rows=1, cols=4)
            tbl3.style = "Table Grid"
            _hdr_row(tbl3, ["Soil", "% Area", "Suitability", "AWC (mm/m)"], bg=HEX_SOLO_HDR)
            _col_widths(tbl3, [7.5, 2.5, 3.0, 2.5])
            for i, st in enumerate(soil_types):
                apt_c  = st.get("apt_class")
                apt_ok = SOIL_APTITUDE_CLASSES.get(apt_c, {}).get("suitable", True)
                apt_s  = f"Cls {apt_c}" if apt_c else "N/D"
                cor    = C_VERDE if apt_ok else C_VERMELHO
                _data_row(tbl3, [
                    f"{st['code']} — {st['name'][:35]}",
                    f"{st['pct_area']:.1f}%",
                    apt_s,
                    str(st["water_props"].get("AWC", "N/D")),
                ], alt=(i % 2 == 1), colors=[None, None, cor, None])
            doc.add_paragraph()

        #soil performance × event
        soil_check = analysis.get("soil_check", {})
        amplifier  = soil_check.get("amplifier", 1.0) if soil_check else 1.0
        if amplifier != 1.0:
            bg_amp = HEX_VERM_BG if amplifier > 1.0 else HEX_VERDE_BG
            cor_amp = C_VERMELHO if amplifier > 1.0 else C_VERDE
            p = doc.add_paragraph()
            _shd(p, bg_amp)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            action = "AMPLIFICA" if amplifier > 1.0 else "ATENUA"
            _run(p, (
                f"  Soil × Event Interaction: soil {action} damage by {amplifier:.2f}x "
                f"(retention: {soil_check.get('retention','N/D')} | "
                f"AWC: {soil_check.get('AWC','N/D')} mm/m)"
            ), bold=True, size=10, color=cor_amp)
        doc.add_paragraph()

    #checklist
    def _checks(self, analysis):
        doc    = self.doc
        checks = [c for c in analysis.get("checks", []) if c.get("weight", 0) > 0]
        doc.add_paragraph()
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Validation Criterion", "Result", "Weight", "OK?"])
        _col_widths(tbl, [6.5, 6.0, 1.5, 1.5])
        for i, chk in enumerate(checks):
            ok  = chk.get("passed", False)
            cor = C_VERDE if ok else C_VERMELHO
            _data_row(tbl, [
                chk.get("name", ""),
                chk.get("value", ""),
                f"{chk.get('weight', 0):.1f}",
                "YES ✔" if ok else "NO ✘",
            ], alt=(i % 2 == 1), colors=[None, None, None, cor])
        doc.add_paragraph()
        sm = analysis.get("summary", {})
        p  = doc.add_paragraph()
        _run(p, (
            f"Criteria: {sm.get('checks_total',0)}  |  "
            f"Aprovados: {sm.get('checks_passed',0)}  |  "
            f"Score: {sm.get('score_raw','N/D')}"
        ), bold=True, size=10, color=C_AZUL_ESC)
        doc.add_paragraph()

    #casualty
    def _loss(self, analysis, hist_baseline: Dict = None):
        loss = analysis.get("loss_estimate", {})
        if not loss:
            return
        doc = self.doc
        doc.add_paragraph()
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        _hdr_row(tbl, ["Economic Indicator", "Value"])
        _col_widths(tbl, [8.5, 7.0])

        fields = [
            ("Total assessed area",          f"{loss.get('area_ha', 0):.1f} ha"),
            ("Expected yield",        f"{loss.get('expected_yield_bags_ha', 0)} bags/ha"),
        ]

        #local history
        hist = loss.get("hist_baseline") or hist_baseline
        if hist and hist.get("local_yield_est_bags_ha"):
            fields.append((
                f"Local historical yield ({hist.get('n_years','?')} years)",
                f"~{hist['local_yield_est_bags_ha']:.1f} sc/ha ({hist.get('years_used','')})",
            ))

        #solo factor
        soil_amp = loss.get("soil_amplifier")
        if soil_amp:
            amp = soil_amp["amplifier"]
            action = "amplified" if amp > 1.0 else "mitigated"
            fields.append((
                f"Fator de solo ({action.lower()})",
                f"{amp:.2f}x — {soil_amp.get('soil_name','N/D')} (AWC {soil_amp.get('AWC','N/D')} mm/m)",
            ))

        fields += [
            ("Estimated actual yield",   f"{loss.get('estimated_yield_bags_ha', 0)} bags/ha"),
            ("Queda de produtividade",        _pct(loss.get("yield_loss_pct", 0))),
            ("Perda total em bags",          f"{loss.get('yield_loss_total_sacas', 0):,.0f} bags"),
            ("Reference price (CEPEA)",      f"{_brl(loss.get('price_brl_bag', 0))}/bag"),
            ("Receita bruta esperada",        _brl(loss.get("expected_revenue_brl", 0))),
            ("Receita estimada com perdas",   _brl(loss.get("estimated_revenue_brl", 0))),
        ]
        for i, (k, v) in enumerate(fields):
            row = tbl.add_row()
            alt = (i % 2 == 1)
            for j, (cell, txt) in enumerate(zip(row.cells, [k, v])):
                if alt:
                    _shd(cell, HEX_ROW_ALT)
                _cell_margins(cell)
                p = cell.paragraphs[0]
                _run(p, txt, bold=(j == 0), size=10)
        #final loss line
        row = tbl.add_row()
        for cell in row.cells:
            _shd(cell, "FFD6D6")
            _cell_margins(cell)
        for j, (cell, txt) in enumerate(zip(row.cells, [
            "PERDA FINANCEIRA ESTIMADA",
            _brl(loss.get("financial_loss_brl", 0)),
        ])):
            p = cell.paragraphs[0]
            _run(p, txt, bold=True, size=11, color=(C_VERMELHO if j == 1 else C_AZUL_ESC))
        doc.add_paragraph()
        comp = loss.get("loss_frac_components", {})
        p = doc.add_paragraph()
        _run(p, (
            f"Components: Climate deficit {comp.get('climate_loss',0):.0f}%  |  "
            f"Anomalia NDVI {comp.get('ndvi_loss',0):.0f}%  |  "
            f"Phenological sensitivity {comp.get('phase_sensitivity',0):.0f}%  |  "
            f"Amplificador solo {comp.get('soil_amplifier',1.0):.2f}x"
        ), size=9, color=C_CINZA, italic=True)
        doc.add_paragraph()

    #verdict
    def _verdict(self, analysis):
        doc     = self.doc
        verdict = analysis.get("verdict", "INCONCLUSIVE")
        conf    = analysis.get("confidence", 0)
        loss    = analysis.get("loss_estimate", {})
        cm = {
            "CONFIRMED":     (C_VERDE,    HEX_VERDE_BG, "✔ CLAIM CONFIRMED"),
            "INCONCLUSIVE":   (C_AMARELO,  HEX_AMAR_BG,  "⚠  INCONCLUSIVE"),
            "NOT CONFIRMED": (C_VERMELHO, HEX_VERM_BG,  "✘ CLAIM NOT CONFIRMED"),
        }
        cor, bg, label = cm.get(verdict, (C_CINZA, "EEEEEE", verdict))
        doc.add_paragraph()
        p = doc.add_paragraph()
        _shd(p, bg)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(6)
        _run(p, label, bold=True, size=18, color=cor)
        p2 = doc.add_paragraph()
        _shd(p2, bg)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after  = Pt(4)
        _run(p2, f"Confidence Level: {conf:.0f}%", bold=True, size=12, color=C_AZUL_ESC)
        #severity
        idw = analysis.get("idw_score", 0)
        sev = analysis.get("severity", "")
        if idw:
            p3 = doc.add_paragraph()
            _shd(p3, bg)
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p3.paragraph_format.space_after = Pt(10)
            _run(p3, f"Severidade: {sev}  |  Score IDW: {idw:.0f}/100",
                 bold=True, size=11, color=C_CINZA)
        doc.add_paragraph()
        el    = EVENT_LABELS.get(self.event_type, self.event_type)
        pi    = analysis.get("phase_info", {})
        phase = pi.get("phase", "desconhecida")
        sens  = pi.get("sensitivity", 0) * 100
        crop  = self.crop_params["name_en"]
        p     = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.8)
        soil_check = analysis.get("soil_check", {})
        soil_note  = ""
        if soil_check and soil_check.get("available"):
            amp = soil_check.get("amplifier", 1.0)
            soil_note = (
                f" Solo {soil_check.get('soil_name','N/D')} "
                f"({'amplificou' if amp > 1.0 else 'atenuou'} o dano em {amp:.2f}x)."
            )
        if verdict == "CONFIRMED":
            _run(p, (
                f"Satellite data (Copernicus/Sentinel-2) and climate records "
                f"(Poseidon Network) present converging evidence of {el}. "
                f"Phenological phase: "
            ))
            _run(p, phase.upper(), bold=True, color=C_AZUL_ESC)
            _run(p, f" (sensibilidade {sens:.0f}%).{soil_note} Perda financeira estimada: ")
            _run(p, _brl(loss.get("financial_loss_brl", 0)), bold=True, color=C_VERMELHO)
            _run(p, f". Confidence: {conf:.0f}%.)
        elif verdict == "INCONCLUSIVE":
            _run(p, (
                f"Partial signals of {el} detected, without sufficient evidence for "
                f"definitive confirmation (confidence: {conf:.0f}%).{soil_note} "
                f"Recomenda-se vistoria presencial."
            ))
        else:
            _run(p, (
                f"Analysed data do not corroborate the claim of {el} "
                f"({conf:.0f}% confidence).{soil_note} Claim not validated."
            ))
        doc.add_paragraph()

    #final note
    def _note(self, soil_data: Dict = None):
        doc = self.doc
        p   = doc.add_paragraph()
        _para_bot_border(p, color="CCCCCC", sz="4")
        solo_str = ""
        if soil_data and not soil_data.get("error"):
            solo_str = " Agricultural suitability: EMBRAPA/IBGE."
        _run(p, (
            "Report automatically generated by the Poseidon-Copernicus-EMBRAPA Validator System. "
            "Dados satelitais: Copernicus/Sentinel-2 (CDSE). "
            "Climate data: Poseidon Network."
            f"{solo_str} "
            "Financial values: CEPEA reference prices. Technical and estimative in nature."
        ), size=8, color=C_CINZA_CLR, italic=True)
